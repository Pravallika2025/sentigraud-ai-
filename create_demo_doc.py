"""
SentinelGPT — Project Demonstration Document
Clean, professional, screenshot-rich, B&W bordered
"""
import os
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from PIL import Image as PILImage, ImageDraw, ImageFont

BASE  = r"c:\Users\User\pravallika sentinel"
IMGS  = os.path.join(BASE, "docs", "images")
DESK  = (r"C:\Users\User\OneDrive\Desktop"
         if os.path.exists(r"C:\Users\User\OneDrive\Desktop")
         else r"C:\Users\User\Desktop")
OUT   = os.path.join(DESK, "SentinelGPT_Demo.docx")

DASH    = os.path.join(IMGS, "dashboard.png")
LOGIN   = os.path.join(IMGS, "login_page.png")
CHAT    = os.path.join(IMGS, "ai_chat.png")
SCAN    = os.path.join(IMGS, "file_scanner.png")

# ─────────────────────────────────────────────────────────────────────────────
#  Generate simple B&W architecture diagram
# ─────────────────────────────────────────────────────────────────────────────
def make_arch():
    W, H = 820, 420
    img = PILImage.new("RGB", (W, H), "#FFFFFF")
    d   = ImageDraw.Draw(img)
    try:
        fb = ImageFont.truetype("arial.ttf", 13)
        fn = ImageFont.truetype("arial.ttf", 11)
    except:
        fb = fn = ImageFont.load_default()

    def box(x1,y1,x2,y2,t,s=""):
        d.rectangle([x1,y1,x2,y2], fill="#FFFFFF", outline="#000000", width=2)
        cx,cy = (x1+x2)//2, (y1+y2)//2
        if s:
            d.text((cx, cy-9), t, fill="#000000", font=fb, anchor="mm")
            d.text((cx, cy+9), s, fill="#555555", font=fn, anchor="mm")
        else:
            d.text((cx, cy),   t, fill="#000000", font=fb, anchor="mm")

    def arr(x1,y1,x2,y2):
        d.line([x1,y1,x2,y2], fill="#000000", width=2)
        d.polygon([(x2,y2),(x2-5,y2-8),(x2+5,y2-8)], fill="#000000")

    box(280,8,540,48, "SECURITY ADMINISTRATOR")
    arr(410,48,410,70)
    box(80,70,740,112, "REACT 19 FRONTEND SPA",
        "Login  |  Dashboard  |  Quarantine Control  |  AI Chat  |  File Scanner")
    arr(410,112,410,138)
    box(100,138,720,180, "FASTAPI BACKEND  (Python 3.11)",
        "REST API  |  JWT Auth  |  WebSocket  |  Heuristic Scoring Engine")
    arr(410,180,410,205)
    d.line([215,205,605,205], fill="#000000", width=2)
    arr(215,205,215,228); arr(605,205,605,228)
    box(30,228,390,300,  "TELEMETRY MONITOR AGENT",
        "Risk Score 0-100  |  MITRE ATT&CK Mapping")
    box(420,228,790,300, "AUTONOMOUS QUARANTINE AGENT",
        "Auto-Block Score>=75  |  AI Remediation")
    d.line([215,300,215,318], fill="#000000", width=2)
    d.line([605,300,605,318], fill="#000000", width=2)
    d.line([215,318,605,318], fill="#000000", width=2)
    arr(410,318,410,340)
    box(140,340,680,382, "SQLITE DATABASE  (SQLAlchemy ORM)",
        "Users  |  SOC Incidents  |  Blocked IPs  |  Perimeter Logs")
    arr(410,382,410,400)
    box(230,400,590,415, "VERCEL CLOUD DEPLOYMENT")

    path = os.path.join(IMGS, "demo_arch.png")
    img.save(path)
    return path

# ─────────────────────────────────────────────────────────────────────────────
#  Page border
# ─────────────────────────────────────────────────────────────────────────────
def page_border(sec):
    sp = sec._sectPr
    pb = OxmlElement("w:pgBorders")
    pb.set(qn("w:offsetFrom"), "page")
    for side in ("top","left","bottom","right"):
        el = OxmlElement(f"w:{side}")
        el.set(qn("w:val"),   "single")
        el.set(qn("w:sz"),    "8")
        el.set(qn("w:space"), "20")
        el.set(qn("w:color"), "000000")
        pb.append(el)
    sp.append(pb)

# ─────────────────────────────────────────────────────────────────────────────
#  Build document
# ─────────────────────────────────────────────────────────────────────────────
def build():
    arch = make_arch()
    doc  = Document()

    sec = doc.sections[0]
    sec.top_margin    = Inches(0.9)
    sec.bottom_margin = Inches(0.9)
    sec.left_margin   = Inches(1.1)
    sec.right_margin  = Inches(1.1)
    page_border(sec)

    sn = doc.styles["Normal"]
    sn.font.name = "Times New Roman"
    sn.font.size = Pt(11)

    # ── helpers ──────────────────────────────────────────────────────────────
    def R(run, bold=False, size=11, italic=False, color=(0,0,0)):
        run.font.name = "Times New Roman"
        run.font.size = Pt(size)
        run.font.bold = bold
        run.font.italic = italic
        run.font.color.rgb = RGBColor(*color)

    def centered(text, size=11, bold=False, italic=False, space_before=0, space_after=6):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(space_before)
        p.paragraph_format.space_after  = Pt(space_after)
        R(p.add_run(text), bold=bold, size=size, italic=italic)

    def section_title(text):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(14)
        p.paragraph_format.space_after  = Pt(5)
        r = p.add_run(text)
        R(r, bold=True, size=14)
        pPr  = p._p.get_or_add_pPr()
        pBdr = OxmlElement("w:pBdr")
        bot  = OxmlElement("w:bottom")
        for k,v in [("w:val","single"),("w:sz","6"),("w:space","1"),("w:color","000000")]:
            bot.set(qn(k),v)
        pBdr.append(bot); pPr.append(pBdr)

    def body(text, bold_label=None, justify=True):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after  = Pt(5)
        p.paragraph_format.line_spacing = Pt(19)
        if justify: p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        if bold_label:
            R(p.add_run(bold_label), bold=True, size=11)
        R(p.add_run(text), size=11)

    def feature_bullet(icon_text, label, desc):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after  = Pt(3)
        p.paragraph_format.left_indent  = Inches(0.15)
        p.paragraph_format.line_spacing = Pt(18)
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        R(p.add_run(f"{icon_text}  {label}  "), bold=True, size=11)
        R(p.add_run(desc), size=11)

    def img_with_caption(path, caption_text, w=5.2):
        if not os.path.exists(path):
            body(f"[Screenshot not found: {os.path.basename(path)}]")
            return
        # image
        pi = doc.add_paragraph()
        pi.alignment = WD_ALIGN_PARAGRAPH.CENTER
        pi.paragraph_format.space_before = Pt(8)
        pi.paragraph_format.space_after  = Pt(2)
        pi.add_run().add_picture(path, width=Inches(w))
        # caption
        pc = doc.add_paragraph()
        pc.alignment = WD_ALIGN_PARAGRAPH.CENTER
        pc.paragraph_format.space_after = Pt(10)
        R(pc.add_run(caption_text), italic=True, size=10)

    # ─────────────────────────────────────────────────────────────────────────
    #  PAGE 1 — TITLE BLOCK
    # ─────────────────────────────────────────────────────────────────────────
    centered("PROJECT DEMONSTRATION", size=12, bold=True, space_before=4, space_after=4)
    centered(
        "SentinelGPT: An AI-Powered Large Language Model\n"
        "Framework for Advanced Cyber Threat Detection and Analysis",
        size=20, bold=True, space_before=0, space_after=10
    )
    centered("_" * 65, size=10, space_before=0, space_after=8)
    centered("pravallika kalangi   |   24VV1F0044   |   MCA 2nd Year   |   July 2026",
             size=12, bold=False, space_before=0, space_after=14)

    # Project overview paragraph
    body(
        "SentinelGPT is a full-stack AI-powered Security Operations Center (SOC) platform that "
        "automates cyber threat detection, real-time network telemetry monitoring, and autonomous "
        "IP quarantine using intelligent software agents. Built with React 19, FastAPI, SQLAlchemy, "
        "and deployed on Vercel Serverless, the system eliminates manual log review, reduces incident "
        "response latency to near zero, and provides 24/7 network perimeter protection.",
        bold_label="Overview:  "
    )

    # Key features — quick-scan row
    section_title("Key Features at a Glance")
    features = [
        ("*", "Real-Time SOC Dashboard:",
         "Live threat velocity charts, risk gauges, severity donuts, and perimeter heatmaps updated continuously via WebSocket."),
        ("*", "Heuristic Threat Scoring:",
         "Autonomous risk scoring engine (0-100) that classifies every network event as Critical, High, Medium, or Low severity."),
        ("*", "MITRE ATT&CK Mapping:",
         "Every detected incident is automatically tagged with the matching MITRE ATT&CK tactic and technique ID."),
        ("*", "Autonomous IP Quarantine:",
         "Intelligent agent automatically isolates any source IP scoring >= 75 — no manual intervention required."),
        ("*", "AI Triage Assistant:",
         "Conversational AI chat interface provides instant remediation guidance for any detected threat."),
        ("*", "Heuristic Payload Scanner:",
         "Upload log files or suspicious payloads for automated static analysis and instant security verdict."),
        ("*", "JWT Operator Authentication:",
         "Secure HS256 token-based authentication with 8-hour expiring sessions and Quick Demo Access bypass."),
        ("*", "JSON Incident Export:",
         "One-click export of complete incident history as structured JSON for SIEM integration."),
    ]
    for icon, lbl, desc in features:
        feature_bullet(icon, lbl, desc)

    doc.add_page_break()

    # ─────────────────────────────────────────────────────────────────────────
    #  PAGE 2 — SYSTEM ARCHITECTURE
    # ─────────────────────────────────────────────────────────────────────────
    section_title("System Architecture")
    body(
        "SentinelGPT uses a modular, decoupled architecture. The operator authenticates via the "
        "React 19 SPA frontend. Telemetry data is processed by two autonomous agents — the "
        "Telemetry Monitor Agent (risk scoring + MITRE mapping) and the Autonomous Quarantine Agent "
        "(IP isolation + validation) — before being stored via SQLAlchemy ORM and rendered on the "
        "live dashboard."
    )
    img_with_caption(arch, "Figure 1: SentinelGPT System Architecture — Component & Agent Flow", w=5.4)

    # Tech stack table
    section_title("Technology Stack")
    tbl = doc.add_table(rows=5, cols=2)
    tbl.style = "Table Grid"
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Header row
    for cell, txt in zip(tbl.rows[0].cells, ["Layer", "Technologies"]):
        tc = cell._tc; tp2 = tc.get_or_add_tcPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:val"),"clear"); shd.set(qn("w:color"),"auto"); shd.set(qn("w:fill"),"000000")
        tp2.append(shd)
        pp = cell.paragraphs[0]; pp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        R(pp.add_run(txt), bold=True, size=11, color=(255,255,255))

    stack = [
        ("Frontend",   "React 19,  Vite 5.4,  Recharts,  Web Audio API,  CSS3 Glassmorphism"),
        ("Backend",    "Python 3.11,  FastAPI,  Uvicorn,  PyJWT,  WebSocket"),
        ("Database",   "SQLAlchemy ORM,  SQLite (dev)  /  Vercel Storage (prod)"),
        ("Deployment", "Vercel Serverless,  GitHub Actions CI/CD,  REST + Swagger Docs"),
    ]
    for i,(layer, tech) in enumerate(stack):
        fill = "F2F2F2" if i%2==0 else "FFFFFF"
        for cell, txt, bold in [(tbl.rows[i+1].cells[0], layer, True),
                                  (tbl.rows[i+1].cells[1], tech,  False)]:
            tc = cell._tc; tp2 = tc.get_or_add_tcPr()
            shd = OxmlElement("w:shd")
            shd.set(qn("w:val"),"clear"); shd.set(qn("w:color"),"auto"); shd.set(qn("w:fill"),fill)
            tp2.append(shd)
            pp = cell.paragraphs[0]
            pp.paragraph_format.space_before = Pt(3)
            pp.paragraph_format.space_after  = Pt(3)
            R(pp.add_run(txt), bold=bold, size=11)

    doc.add_page_break()

    # ─────────────────────────────────────────────────────────────────────────
    #  PAGE 3 — DASHBOARD & LOGIN
    # ─────────────────────────────────────────────────────────────────────────
    section_title("1.  Real-Time SOC Operations Dashboard")
    body(
        "The main dashboard provides a comprehensive view of the network security posture. It displays "
        "live KPI cards showing total incidents, active threats, quarantined IPs, and system health. "
        "Recharts-powered area charts stream threat velocity in real time via WebSocket. A risk "
        "distribution donut shows Critical / High / Medium / Low severity breakdown. A perimeter "
        "heatmap reveals the most targeted network sectors. The quarantine panel lists all blocked IPs "
        "with timestamps, MITRE tags, and one-click revocation controls."
    )
    img_with_caption(DASH, "Figure 2: SentinelGPT Real-Time SOC Operations Dashboard", w=5.3)

    section_title("2.  Operator Authentication Portal")
    body(
        "The authentication portal provides secure multi-role login using HS256 JWT tokens with "
        "8-hour session expiry. A Quick Demo Access button allows instant evaluation bypass for "
        "reviewers and evaluators. SHA-256 hashed passwords are stored in the Users table managed "
        "by SQLAlchemy ORM."
    )
    img_with_caption(LOGIN, "Figure 3: Operator Authentication Portal with Quick Demo Access", w=4.8)

    doc.add_page_break()

    # ─────────────────────────────────────────────────────────────────────────
    #  PAGE 4 — AI CHAT & FILE SCANNER
    # ─────────────────────────────────────────────────────────────────────────
    section_title("3.  AI Threat Triage Assistant")
    body(
        "The conversational AI Triage Assistant is integrated directly into the dashboard. Security "
        "operators can describe a threat scenario or paste incident details and receive instant, "
        "structured remediation guidance. The assistant understands MITRE ATT&CK context, suggests "
        "specific firewall rules, patch priorities, and incident escalation steps — dramatically "
        "reducing mean time to respond (MTTR) for SOC teams."
    )
    img_with_caption(CHAT, "Figure 4: Conversational AI Threat Triage Assistant Interface", w=4.8)

    section_title("4.  Heuristic Payload & Log File Scanner")
    body(
        "The built-in Payload Scanner allows administrators to upload suspicious log files, binary "
        "payloads, or configuration files for automated static analysis. The heuristic engine inspects "
        "file signatures, entropy patterns, and embedded string patterns to generate an instant "
        "security verdict — SAFE, SUSPICIOUS, or MALICIOUS — with a detailed findings report, "
        "eliminating the need for external third-party analysis tools."
    )
    img_with_caption(SCAN, "Figure 5: Heuristic Payload & Static Log File Scanner", w=4.8)

    doc.add_page_break()

    # ─────────────────────────────────────────────────────────────────────────
    #  PAGE 5 — DEMO FLOW + LINKS
    # ─────────────────────────────────────────────────────────────────────────
    section_title("How to Access & Demonstrate the System")
    body("Follow these steps to run a live demonstration of SentinelGPT:", bold_label="Steps:  ")

    steps = [
        ("Step 1 — Open the Live Dashboard",
         'Visit https://sentinelgpt-ai.vercel.app in any web browser. The login screen will appear automatically.'),
        ("Step 2 — Login",
         'Click the "Quick Demo Access" button for instant entry — OR — use credentials: admin / admin123'),
        ("Step 3 — View Real-Time Telemetry",
         'The SOC dashboard loads with live KPI cards, threat velocity charts, severity donuts, and the quarantine list.'),
        ("Step 4 — Simulate a Threat",
         'Click "Simulate Threat" to inject a test attack event. Watch the risk score appear on the chart in real time.'),
        ("Step 5 — Auto-Quarantine",
         'If the simulated threat scores >= 75 (Critical), the Autonomous Agent immediately adds the IP to the Blocked IPs list.'),
        ("Step 6 — AI Triage Chat",
         'Navigate to the AI Chat tab. Ask the assistant about the detected threat for instant remediation guidance.'),
        ("Step 7 — File Scanner",
         'Navigate to the Scanner tab. Upload any .log or .txt file to receive a heuristic security verdict instantly.'),
        ("Step 8 — JSON Export",
         'Click Export to download the full incident history as a structured JSON file for SIEM review.'),
    ]
    for step_lbl, step_text in steps:
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after  = Pt(4)
        p.paragraph_format.left_indent  = Inches(0.15)
        p.paragraph_format.line_spacing = Pt(18)
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        R(p.add_run(f"{step_lbl}:  "), bold=True, size=11)
        R(p.add_run(step_text), size=11)

    # Deployment Links Table
    section_title("Live Deployment & GitHub Links")
    body("The following official links provide access to the deployed application, API documentation, "
         "and complete source code:")

    lt = doc.add_table(rows=3, cols=2)
    lt.style = "Table Grid"
    lt.alignment = WD_TABLE_ALIGNMENT.CENTER
    link_data = [
        ("Live SOC Dashboard",           "https://sentinelgpt-ai.vercel.app"),
        ("Swagger API Documentation",    "https://sentinelgpt-ai.vercel.app/docs"),
        ("GitHub Project Repository",    "https://github.com/Pravallika2025/sentigraud-ai-.git"),
    ]
    for i,(lbl,url) in enumerate(link_data):
        fill = "F2F2F2" if i%2==0 else "FFFFFF"
        for cell, txt, bold in [(lt.rows[i].cells[0], lbl, True),
                                  (lt.rows[i].cells[1], url, False)]:
            tc = cell._tc; tp2 = tc.get_or_add_tcPr()
            shd = OxmlElement("w:shd")
            shd.set(qn("w:val"),"clear"); shd.set(qn("w:color"),"auto"); shd.set(qn("w:fill"),fill)
            tp2.append(shd)
            pp = cell.paragraphs[0]
            pp.paragraph_format.space_before = Pt(4)
            pp.paragraph_format.space_after  = Pt(4)
            R(pp.add_run(txt), bold=bold, size=11)

    # Footer
    doc.add_paragraph()
    fp = doc.add_paragraph()
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    fp.paragraph_format.space_before = Pt(12)
    R(fp.add_run(
        "━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━\n"
        "SentinelGPT  |  pravallika kalangi  |  24VV1F0044  |  MCA 2nd Year  |  2026"
    ), italic=True, size=10)

    # ── Save & open ───────────────────────────────────────────────────────────
    doc.save(OUT)
    sz = round(os.path.getsize(OUT)/1024)
    print(f"Saved ({sz} KB): {OUT}")
    import os as _os; _os.startfile(OUT)
    print("Opened on screen!")

if __name__ == "__main__":
    build()
