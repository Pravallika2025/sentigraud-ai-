"""
Professional Word Document Generator for SentinelGPT Project Report
- Times New Roman font throughout
- 4-sided page border (black)
- Proper line spacing (1.5) and word spacing
- Black & White professional presentation
- Embedded architecture and workflow diagrams
- Exact section order as per reference image
"""
import os
import sys
from docx import Document
from docx.shared import Inches, Pt, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml, OxmlElement
from PIL import Image, ImageDraw, ImageFont

BASE_DIR    = r"c:\Users\User\pravallika sentinel"
IMG_DIR     = os.path.join(BASE_DIR, "docs", "images")
DESKTOP     = (r"C:\Users\User\OneDrive\Desktop"
               if os.path.exists(r"C:\Users\User\OneDrive\Desktop")
               else r"C:\Users\User\Desktop")
os.makedirs(IMG_DIR, exist_ok=True)

OUT_FILE    = os.path.join(DESKTOP, "SentinelGPT_Final_Report.docx")
DOCS_FILE   = os.path.join(BASE_DIR, "docs", "SentinelGPT_Final_Report.docx")

# ─────────────────────────────────────────────────────────────────────────────
# FONT HELPERS
# ─────────────────────────────────────────────────────────────────────────────
def _fnt(name, size):
    for f in [name, "times.ttf", "timesnewroman.ttf"]:
        try:
            return ImageFont.truetype(f, size)
        except:
            pass
    return ImageFont.load_default()

# ─────────────────────────────────────────────────────────────────────────────
# 1.  GENERATE: System Architecture Diagram (B&W, clean lines)
# ─────────────────────────────────────────────────────────────────────────────
def make_arch_diagram():
    W, H = 800, 560
    img  = Image.new("RGB", (W, H), "#FFFFFF")
    d    = ImageDraw.Draw(img)
    fb   = _fnt("timesbd.ttf", 13)
    fn   = _fnt("times.ttf",   11)

    def box(x1, y1, x2, y2, top, sub=""):
        d.rectangle([x1,y1,x2,y2], fill="#FFFFFF", outline="#000000", width=2)
        cx = (x1+x2)//2
        cy = (y1+y2)//2 - 8 if sub else (y1+y2)//2
        d.text((cx, cy),      top, fill="#000000", font=fb, anchor="mm")
        if sub:
            d.text((cx, cy+17), sub, fill="#333333", font=fn, anchor="mm")

    def arrow(x1,y1,x2,y2):
        d.line([x1,y1,x2,y2], fill="#000000", width=2)
        d.polygon([(x2,y2),(x2-6,y2-10),(x2+6,y2-10)], fill="#000000")

    box(260,12,540,52,  "SECURITY ADMINISTRATOR","( Authentication & Dashboard Access )")
    arrow(400,52,400,77)
    box(80,77,720,127,  "FRONTEND INTERFACE  —  React 19 + Vite 5.4  SPA",
        "Login  |  SOC Dashboard  |  Telemetry Feed  |  AI Chat  |  File Scanner")
    arrow(400,127,400,152)
    box(100,152,700,202,"FASTAPI BACKEND CONTROLLER  (index.py / main.py)",
        "REST Router  |  JWT Authentication  |  CORS Middleware  |  WebSocket Engine")
    arrow(400,202,400,227)
    d.line([220,227,580,227], fill="#000000", width=2)
    arrow(220,227,220,252)
    arrow(580,227,580,252)
    box(30,252,385,337, "TELEMETRY MONITOR AGENT",
        "Heuristic Risk Scoring  (0 - 100)\nMITRE ATT&CK Technique Mapping\nVelocity Anomaly Detection")
    box(415,252,770,337,"AUTONOMOUS QUARANTINE AGENT",
        "Auto-Block Score >= 75\nAI Remediation Guidance\nManual Revoke Validation")
    d.line([220,337,220,362], fill="#000000", width=2)
    d.line([580,337,580,362], fill="#000000", width=2)
    d.line([220,362,580,362], fill="#000000", width=2)
    arrow(400,362,400,387)
    box(130,387,670,437,"SQLALCHEMY ORM  —  SQLITE DATABASE",
        "Users  |  SOC Incidents  |  Blocked IPs  |  Perimeter Logs")
    arrow(400,437,400,462)
    box(100,462,700,512,"VERCEL SERVERLESS CLOUD DEPLOYMENT",
        "Live SOC Dashboard  |  JSON Export  |  Swagger API Docs")

    path = os.path.join(IMG_DIR, "final_arch.png")
    img.save(path)
    return path

# ─────────────────────────────────────────────────────────────────────────────
# 2.  GENERATE: Technology Stack / Agent Workflow Diagram
# ─────────────────────────────────────────────────────────────────────────────
def make_flow_diagram():
    W, H = 800, 520
    img  = Image.new("RGB", (W, H), "#FFFFFF")
    d    = ImageDraw.Draw(img)
    fb   = _fnt("timesbd.ttf", 13)
    fn   = _fnt("times.ttf",   11)

    def box(x1,y1,x2,y2,top,sub=""):
        d.rectangle([x1,y1,x2,y2], fill="#FFFFFF", outline="#000000", width=2)
        cx=(x1+x2)//2; cy=(y1+y2)//2-8 if sub else (y1+y2)//2
        d.text((cx,cy),     top, fill="#000000", font=fb, anchor="mm")
        if sub:
            d.text((cx,cy+17),sub, fill="#333333", font=fn, anchor="mm")

    def arrow(x1,y1,x2,y2):
        d.line([x1,y1,x2,y2], fill="#000000", width=2)
        d.polygon([(x2,y2),(x2-6,y2-10),(x2+6,y2-10)], fill="#000000")

    rows = [
        (80,10,720,55,  "NETWORK TRAFFIC & SECURITY TELEMETRY FEED",
                        "Inbound Logs  |  IP Probes  |  HTTP Requests  |  Packet Streams"),
        (80,85,720,135, "FRONTEND DASHBOARD & ALERT FEED",
                        "Real-Time Metrics  |  Operator Controls  |  Telemetry Visualization"),
        (80,165,720,225,"TELEMETRY MONITOR & THREAT DETECTION AGENT",
                        "IP Velocity Analysis  |  Risk Score (0-100)  |  MITRE ATT&CK Mapping"),
        (80,255,720,315,"AUTONOMOUS QUARANTINE & VALIDATION AGENT",
                        "Threshold Enforcement (>=75)  |  AI Remediation  |  False-Positive Prevention"),
        (160,345,640,395,"DATABASE STORAGE LAYER",
                         "SOC Incidents  |  Blocked IPs  |  User Accounts  |  Perimeter Logs"),
    ]
    for i,(x1,y1,x2,y2,t,s) in enumerate(rows):
        box(x1,y1,x2,y2,t,s)
        if i < len(rows)-1:
            arrow(400, y2, 400, y2+30)

    arrow(400,395,400,420)
    # Output boxes row
    d.line([80,420,720,420], fill="#000000", width=2)
    outs = [("DASHBOARD",120),("CHARTS",255),("QUARANTINE",390),
            ("AI CHAT",525),("SCANNER",660)]
    for lbl,cx in outs:
        arrow(cx,420,cx,435)
        box(cx-60,435,cx+60,480,lbl)

    path = os.path.join(IMG_DIR, "final_flow.png")
    img.save(path)
    return path

# ─────────────────────────────────────────────────────────────────────────────
# 3.  PAGE BORDER HELPER  (4-sided 1pt black border on every page)
# ─────────────────────────────────────────────────────────────────────────────
def add_page_border(section):
    """Add a thin black border to every page via pgBorders XML."""
    sectPr = section._sectPr
    pgBdr  = OxmlElement("w:pgBorders")
    pgBdr.set(qn("w:offsetFrom"), "page")
    for side in ("top","left","bottom","right"):
        el = OxmlElement(f"w:{side}")
        el.set(qn("w:val"),   "single")
        el.set(qn("w:sz"),    "4")        # 0.5pt
        el.set(qn("w:space"), "24")
        el.set(qn("w:color"), "000000")
        pgBdr.append(el)
    sectPr.append(pgBdr)

# ─────────────────────────────────────────────────────────────────────────────
# 4.  BUILD THE DOCUMENT
# ─────────────────────────────────────────────────────────────────────────────
def build(save_paths):
    arch_img = make_arch_diagram()
    flow_img = make_flow_diagram()

    doc = Document()

    # ── Margins ──────────────────────────────────────────────────────────────
    sec = doc.sections[0]
    sec.top_margin    = Inches(1.0)
    sec.bottom_margin = Inches(1.0)
    sec.left_margin   = Inches(1.25)
    sec.right_margin  = Inches(1.25)
    add_page_border(sec)

    # ── Default paragraph style ──────────────────────────────────────────────
    ns = doc.styles["Normal"]
    ns.font.name           = "Times New Roman"
    ns.font.size           = Pt(12)
    ns.font.color.rgb      = RGBColor(0,0,0)
    ns.paragraph_format.space_before  = Pt(0)
    ns.paragraph_format.space_after   = Pt(6)
    ns.paragraph_format.line_spacing  = Pt(18)   # ~1.5 of 12pt

    # ═══════════════════════════════════════════════════════════════
    #  STYLE HELPERS
    # ═══════════════════════════════════════════════════════════════
    def set_run(run, bold=False, size=12, italic=False):
        run.font.name      = "Times New Roman"
        run.font.size      = Pt(size)
        run.font.bold      = bold
        run.font.italic    = italic
        run.font.color.rgb = RGBColor(0,0,0)

    def h1(text):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(14)
        p.paragraph_format.space_after  = Pt(6)
        p.paragraph_format.keep_with_next = True
        r = p.add_run(text)
        set_run(r, bold=True, size=15)
        # underline rule
        pPr  = p._p.get_or_add_pPr()
        pBdr = OxmlElement("w:pBdr")
        bot  = OxmlElement("w:bottom")
        bot.set(qn("w:val"),   "single")
        bot.set(qn("w:sz"),    "6")
        bot.set(qn("w:space"), "2")
        bot.set(qn("w:color"), "000000")
        pBdr.append(bot)
        pPr.append(pBdr)

    def h2(text):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(10)
        p.paragraph_format.space_after  = Pt(3)
        r = p.add_run(text)
        set_run(r, bold=True, size=13)

    def body(text, justify=True, bold_label=None):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after  = Pt(6)
        p.paragraph_format.line_spacing = Pt(20)  # 1.5 spacing for 12pt
        if justify:
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        if bold_label:
            r1 = p.add_run(bold_label)
            set_run(r1, bold=True, size=12)
        r2 = p.add_run(text)
        set_run(r2, size=12)

    def caption(text):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(2)
        p.paragraph_format.space_after  = Pt(10)
        r = p.add_run(text)
        set_run(r, italic=True, size=10)

    def img_block(path, w_inches=5.2, cap=None):
        if not os.path.exists(path):
            return
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(10)
        p.paragraph_format.space_after  = Pt(4)
        p.add_run().add_picture(path, width=Inches(w_inches))
        if cap:
            caption(cap)

    def styled_table_hdr(cell, text):
        """Black background, white bold text."""
        tc  = cell._tc
        tcPr= tc.get_or_add_tcPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:val"),   "clear")
        shd.set(qn("w:color"), "auto")
        shd.set(qn("w:fill"),  "000000")
        tcPr.append(shd)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(text)
        set_run(r, bold=True, size=11)
        r.font.color.rgb = RGBColor(0xFF,0xFF,0xFF)

    def alt_row(cell, shade=False):
        if shade:
            tc  = cell._tc
            tcPr= tc.get_or_add_tcPr()
            shd = OxmlElement("w:shd")
            shd.set(qn("w:val"),   "clear")
            shd.set(qn("w:color"), "auto")
            shd.set(qn("w:fill"),  "F2F2F2")
            tcPr.append(shd)

    # ═══════════════════════════════════════════════════════════════
    #  PAGE 1  —  COVER PAGE
    # ═══════════════════════════════════════════════════════════════
    # top spacer
    for _ in range(6):
        doc.add_paragraph()

    # Project title (centered, bold, 22pt)
    tp = doc.add_paragraph()
    tp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    tp.paragraph_format.space_after = Pt(10)
    tr = tp.add_run(
        "SentinelGPT: An AI-Powered Large Language Model\n"
        "Framework for Advanced Cyber Threat Detection\n"
        "and Analysis"
    )
    set_run(tr, bold=True, size=22)

    # Horizontal rule
    hr = doc.add_paragraph()
    hr.alignment = WD_ALIGN_PARAGRAPH.CENTER
    hr.paragraph_format.space_after = Pt(4)
    hrr = hr.add_run("─" * 55)
    set_run(hrr, size=12)

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub.paragraph_format.space_after = Pt(60)
    sr = sub.add_run(
        "A Project Report Submitted for the Partial Fulfillment of the\n"
        "Requirements for the Award of the Degree of\n"
        "Master of Computer Applications (MCA)"
    )
    set_run(sr, italic=True, size=13)

    # Submitted by block (right-aligned)
    by_p = doc.add_paragraph()
    by_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    by_p.paragraph_format.space_after = Pt(4)
    by_r = by_p.add_run("Submitted by")
    set_run(by_r, size=12)

    nm_p = doc.add_paragraph()
    nm_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    nm_p.paragraph_format.space_after = Pt(3)
    nm_r = nm_p.add_run("pravallika kalangi")
    set_run(nm_r, bold=True, size=16)

    rn_p = doc.add_paragraph()
    rn_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    rn_p.paragraph_format.space_after = Pt(3)
    rn_r = rn_p.add_run("24VV1F0044")
    set_run(rn_r, size=13)

    yr_p = doc.add_paragraph()
    yr_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    yr_p.paragraph_format.space_after = Pt(3)
    yr_r = yr_p.add_run("MCA — 2nd Year  |  July 2026")
    set_run(yr_r, size=12)

    doc.add_page_break()

    # ═══════════════════════════════════════════════════════════════
    #  PAGE 2  —  PROBLEM STATEMENT
    # ═══════════════════════════════════════════════════════════════
    h1("1.   Problem Statement")
    body(
        "Modern enterprise cybersecurity operations face severe and escalating challenges in managing "
        "the overwhelming volume of network telemetry and security alerts generated by large-scale IT "
        "infrastructure. Security Information and Event Management (SIEM) systems process millions of "
        "event logs, network probes, and API requests every single day across complex distributed cloud "
        "environments. Despite the scale of the threat landscape, many organizations continue to rely "
        "heavily on manual log inspection, static firewall rules, and fragmented monitoring tools. "
        "This outdated approach creates significant operational bottlenecks, leads to severe analyst "
        "alert fatigue, and drastically increases incident response times — allowing malicious cyber "
        "threats to propagate unchecked through internal networks."
    )
    body(
        "One of the most critical challenges in contemporary cyber defense is identifying sophisticated "
        "multi-stage attacks such as credential stuffing, Distributed Denial of Service (DDoS) vectors, "
        "SQL injection payloads, and brute-force authentication spikes. Traditional rule-based intrusion "
        "detection systems generate extremely high rates of false positives while failing to correlate "
        "anomalous network behavior with established attack frameworks like MITRE ATT&CK. Furthermore, "
        "when a critical threat is detected, manual quarantine procedures require administrators to "
        "manually update firewall rules and isolate compromised IP addresses — a slow, delayed process "
        "that exposes internal networks to lateral movement and unauthorized data exfiltration."
    )
    body(
        "Existing security dashboards also lack transparent decision-making explanations, intelligent "
        "payload analysis, and real-time visualization capabilities. Security operators are frequently "
        "left without instant remediation guidance, making it difficult to understand why a specific "
        "threat score was assigned or what immediate countermeasures should be taken."
    )
    body(
        "To address these critical challenges, this project presents SentinelGPT — an AI-Powered "
        "Large Language Model Framework for Advanced Cyber Threat Detection and Analysis. The system "
        "leverages intelligent autonomous software agents to automatically monitor network telemetry, "
        "calculate risk scores using heuristic anomaly algorithms, dynamically map threats to MITRE "
        "ATT&CK techniques, and autonomously quarantine high-risk IP addresses. By integrating "
        "real-time telemetry streaming, agentic quarantine execution, interactive chart analytics, "
        "a conversational AI triage assistant, and static payload scanning into a unified serverless "
        "cloud deployment, SentinelGPT significantly minimizes manual analyst effort, eliminates "
        "incident response latency, and delivers a scalable, reliable solution for modern enterprise "
        "cyber defense operations."
    )
    doc.add_page_break()

    # ═══════════════════════════════════════════════════════════════
    #  PAGE 3  —  DIFFERENCE TABLE
    # ═══════════════════════════════════════════════════════════════
    h1("2.   Difference: Traditional System vs. Proposed SentinelGPT")
    body(
        "The following table highlights the key operational differences between the conventional "
        "manual security management approach and the proposed AI-powered SentinelGPT framework:"
    )

    tbl = doc.add_table(rows=9, cols=2)
    tbl.style = "Table Grid"
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    # Header
    styled_table_hdr(tbl.rows[0].cells[0], "Traditional Security Operations")
    styled_table_hdr(tbl.rows[0].cells[1], "Proposed SentinelGPT System")

    rows_data = [
        ("Security alerts monitored manually using static firewall rules and basic log viewers.",
         "Threats detected and analyzed automatically 24/7 via intelligent software agents."),
        ("High analyst alert fatigue; delayed incident detection and slow response times.",
         "Real-time heuristic anomaly scoring (0-100) eliminates fatigue and detects threats instantly."),
        ("IP quarantine requires manual admin intervention and firewall rule configuration.",
         "Autonomous Quarantine Agent auto-isolates malicious IPs (score >= 75) in real time."),
        ("No standardized MITRE ATT&CK threat framework alignment or technique tagging.",
         "All incidents automatically mapped to MITRE ATT&CK tactics and techniques."),
        ("Threat triage requires lengthy manual analyst research and verification.",
         "Integrated AI Triage Assistant delivers instant remediation guidance and threat context."),
        ("Limited or no real-time visualization of threat velocity and perimeter metrics.",
         "Glassmorphic dashboard with live telemetry charts, risk gauges, and perimeter heatmaps."),
        ("Payload and log file analysis performed manually using external third-party tools.",
         "Built-in Payload Scanner automatically analyzes files with heuristic verdict reports."),
        ("Incident reporting requires manual log gathering and formatting effort.",
         "Instant JSON incident-log export via REST API (/api/export) — one click."),
    ]

    for i, (l, r) in enumerate(rows_data):
        shade = (i % 2 == 0)
        cl, cr = tbl.rows[i+1].cells
        alt_row(cl, shade); alt_row(cr, shade)
        pl = cl.paragraphs[0]
        pl.paragraph_format.space_before = Pt(3)
        pl.paragraph_format.space_after  = Pt(3)
        pl.paragraph_format.line_spacing = Pt(16)
        rl = pl.add_run(l); set_run(rl, size=10)

        pr = cr.paragraphs[0]
        pr.paragraph_format.space_before = Pt(3)
        pr.paragraph_format.space_after  = Pt(3)
        pr.paragraph_format.line_spacing = Pt(16)
        rr = pr.add_run(r); set_run(rr, size=10)

    doc.add_page_break()

    # ═══════════════════════════════════════════════════════════════
    #  PAGE 4  —  PROPOSED SOLUTION  (300 words)
    # ═══════════════════════════════════════════════════════════════
    h1("3.   Proposed Solution")
    body(
        "SentinelGPT is an intelligent, full-stack, cloud-native cyber defense platform designed to "
        "automate threat detection, security monitoring, and incident mitigation for enterprise "
        "environments. Unlike traditional reactive security tools, the system uses autonomous "
        "software agents that continuously evaluate incoming network telemetry against heuristic "
        "scoring models and enforce security policies — replacing slow, error-prone manual workflows "
        "with instantaneous automated decisions."
    )
    body(
        "The Telemetry Monitor Agent serves as the primary threat intelligence engine. It continuously "
        "analyzes packet rates, IP origins, payload signatures, and request patterns to compute a "
        "normalized risk score ranging from 0 (benign) to 100 (critical). Each detected incident is "
        "dynamically mapped to a corresponding MITRE ATT&CK technique — for example, T1078 (Valid "
        "Accounts), T1498 (Network DoS), or T1190 (Exploit Public-Facing Application) — and enriched "
        "with AI-generated remediation guidance specific to the attack pattern."
    )
    body(
        "The Autonomous Quarantine Agent operates independently to enforce firewall isolation policies. "
        "When an incident risk score reaches or exceeds the critical threshold of 75, the agent "
        "immediately writes the offending IP address to the Blocked IPs database table, effectively "
        "quarantining it from all further system access. Security administrators retain complete manual "
        "override capability — they can revoke quarantine entries, trigger simulated threat injections "
        "for testing, or clear incident logs directly from the interactive dashboard."
    )
    body(
        "By unifying intelligent threat detection, automated quarantine enforcement, real-time "
        "telemetry visualization, conversational AI triage, and cloud-native serverless deployment "
        "into a single cohesive platform, SentinelGPT establishes a modern, scalable, and highly "
        "reliable solution for next-generation enterprise cybersecurity operations — eliminating "
        "response latency, reducing manual effort, and maintaining continuous 24/7 network "
        "perimeter protection."
    )

    h1("4.   Overall System Architecture")
    body(
        "SentinelGPT follows a modular, decoupled full-stack architecture. The security operator "
        "authenticates via the React 19 SPA frontend, which communicates with the FastAPI backend "
        "controller for telemetry snapshots and WebSocket streaming. Two autonomous agents process "
        "all incoming security events. All incident logs, quarantine records, and user credentials "
        "are persisted via SQLAlchemy ORM to SQLite (local) or Vercel cloud storage (production). "
        "The complete end-to-end architecture is shown in the diagram below:"
    )
    img_block(arch_img, w_inches=5.2, cap="Figure 1: SentinelGPT End-to-End System Architecture Diagram")
    doc.add_page_break()

    # ═══════════════════════════════════════════════════════════════
    #  PAGE 5  —  TECHNOLOGIES USED
    # ═══════════════════════════════════════════════════════════════
    h1("5.   Technologies Used")
    body(
        "SentinelGPT is built on a carefully selected combination of modern frontend, backend, "
        "database, and AI technologies that together enable automated threat monitoring, real-time "
        "data streaming, and secure cloud deployment:"
    )

    h2("Frontend Stack")
    body("Modern UI component library for building dynamic, reactive glassmorphic dashboard trees "
         "and managing real-time telemetry state updates efficiently.", bold_label="React 19:  ")
    body("Next-generation build tool providing ultra-fast Hot Module Replacement (HMR) and optimized "
         "production bundles for serverless deployment.", bold_label="Vite 5.4:  ")
    body("Data visualization library used to render live threat velocity area charts, "
         "heuristic risk gauges, severity donut charts, and perimeter heatmaps.", bold_label="Recharts:  ")
    body("Browser-native alarm engine that fires audio alerts on Critical/High threat "
         "detection events.", bold_label="Web Audio API:  ")

    h2("Backend Stack")
    body("Core language for implementing agent logic, threat heuristic algorithms, database ORM "
         "models, and all server-side processing.", bold_label="Python 3.11+:  ")
    body("High-performance async web framework for RESTful API routing, WebSocket streaming, "
         "JWT authentication, and automatic Swagger/OpenAPI documentation generation.", bold_label="FastAPI:  ")
    body("ASGI server powering local development execution and Vercel serverless function "
         "routing.", bold_label="Uvicorn:  ")
    body("HS256 JSON Web Token generation and validation providing 8-hour expiring secure "
         "operator sessions.", bold_label="PyJWT:  ")

    h2("Database & Deployment")
    body("Python Object-Relational Mapper providing type-safe, secure database interactions "
         "without writing raw SQL.", bold_label="SQLAlchemy ORM:  ")
    body("Zero-configuration relational database for local development incident logging, "
         "quarantine records, and user accounts.", bold_label="SQLite:  ")
    body("Serverless cloud platform serving the React static build and executing Python "
         "API functions at the network edge.", bold_label="Vercel:  ")
    body("Automated build and deployment pipeline triggered on every push to the main "
         "branch.", bold_label="GitHub Actions CI/CD:  ")

    doc.add_page_break()

    # ═══════════════════════════════════════════════════════════════
    #  PAGE 6  —  ARCHITECTURE + AGENTS + FLOW DIAGRAM
    # ═══════════════════════════════════════════════════════════════
    h1("6.   Agents Used & Architecture Flow")
    img_block(flow_img, w_inches=5.2, cap="Figure 2: Agent Workflow — From Telemetry Ingestion to SOC Output")

    h2("Agent 1  —  Telemetry Monitor & Threat Detection Agent")
    body(
        "Continuously inspects incoming network log events, extracts source IP details and payload "
        "metadata, computes a heuristic risk score (0–100), and assigns a MITRE ATT&CK tactic and "
        "technique tag. Produces a structured SOC Incident payload that is persisted to the database "
        "and streamed live to the frontend dashboard in real time via WebSocket."
    )

    h2("Agent 2  —  Autonomous Quarantine & Validation Agent")
    body(
        "Evaluates every incident against the critical severity threshold (score >= 75). Qualifying "
        "incidents are automatically written to the Blocked IPs table, quarantining the offending "
        "source IP from all further system access. The agent also validates manual revoke requests "
        "from administrators and generates contextual AI-driven remediation advice to prevent "
        "false-positive lockouts."
    )

    h1("7.   Algorithms Used")
    body("Evaluates request velocity, payload anomaly patterns, and IP reputation to produce a "
         "normalized 0–100 risk score — without relying solely on static signatures.",
         bold_label="1.  Heuristic Threat Scoring Algorithm:  ")
    body("Maps computed scores to four priority tiers: Critical (>=75), High (60–74), Medium (40–59), "
         "Low (<40) — directly driving automated quarantine and alert logic.",
         bold_label="2.  Severity Matrix Classifier:  ")
    body("Monitors burst traffic spikes per source IP over sliding time windows to identify active "
         "DDoS attempts and port-scan sweeps in real time.",
         bold_label="3.  Velocity Rate Anomaly Detector:  ")

    doc.add_page_break()

    # ═══════════════════════════════════════════════════════════════
    #  PAGE 7  —  IMPLEMENTATION PHASES
    # ═══════════════════════════════════════════════════════════════
    h1("8.   Implementation Phases")
    phases = [
        ("Phase 1 — Requirement Analysis",
         "Identified SOC pain points, defined threat scoring metrics, mapped MITRE ATT&CK coverage requirements, and finalized functional and non-functional system specifications."),
        ("Phase 2 — System Design",
         "Architected the decoupled React SPA + FastAPI serverless backend; finalized REST API contracts, database schema design, and agent workflow blueprints."),
        ("Phase 3 — Database Schema Design",
         "Designed SQLAlchemy ORM models — Users, SOC Incidents, Perimeter Logs, and Blocked IPs — with SQLite for local development and Vercel storage for production."),
        ("Phase 4 — Frontend Development",
         "Built glassmorphic React components including ThreatChart, RiskGaugeChart, BlockedIPs manager, AI Chat assistant panel, File Scanner interface, and the main Login portal."),
        ("Phase 5 — Backend API & JWT Auth",
         "Implemented all REST routes (/api/snapshot, /api/block_ip, /api/unblock_ip, /api/sim_threat, /api/export, /api/login) with HS256 JWT authentication middleware and CORS configuration."),
        ("Phase 6 — Agentic AI Development",
         "Programmed the Telemetry Monitor Agent (heuristic scoring + MITRE mapping) and the Autonomous Quarantine Agent (threshold enforcement + AI remediation generation)."),
        ("Phase 7 — Integration & Validation Testing",
         "Injected simulated threat payloads to verify scoring accuracy, quarantine trigger reliability, manual override controls, and WebSocket streaming performance under load."),
        ("Phase 8 — Cloud Deployment",
         "Deployed the complete full-stack application to Vercel Serverless; configured environment variables, verified Swagger API docs in production, and tested the live dashboard end-to-end."),
    ]
    for lbl, desc in phases:
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after  = Pt(5)
        p.paragraph_format.line_spacing = Pt(20)
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        rb = p.add_run(f"{lbl}:  ")
        set_run(rb, bold=True, size=12)
        rt = p.add_run(desc)
        set_run(rt, size=12)

    doc.add_page_break()

    # ═══════════════════════════════════════════════════════════════
    #  PAGE 8  —  APPLICATIONS + CONCLUSION
    # ═══════════════════════════════════════════════════════════════
    h1("9.   Applications")
    apps = [
        ("Enterprise SOC Operations",
         "Continuous 24/7 network perimeter monitoring, automated threat triage, and instant IP quarantine for dedicated corporate security teams."),
        ("Financial Institutions",
         "Protects banking portals and payment gateways from credential stuffing, brute-force authentication attacks, and transaction fraud."),
        ("Cloud Service Providers",
         "Safeguards serverless APIs, microservices, and distributed cloud infrastructure from unauthorized probes and zero-day exploits."),
        ("Healthcare Networks",
         "Secures hospital patient data endpoints and clinical systems against ransomware, malware payload injection, and PHI data exfiltration."),
        ("E-Commerce Platforms",
         "Prevents DDoS disruptions and SQL injection attacks during high-traffic sales events, ensuring continuous uptime and transaction security."),
        ("Academic & Research Institutions",
         "Monitors university campus networks, protects research data, and prevents faculty and student account hijacking or credential theft."),
    ]
    for lbl, desc in apps:
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after  = Pt(5)
        p.paragraph_format.line_spacing = Pt(20)
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        rb = p.add_run(f"{lbl}:  ")
        set_run(rb, bold=True, size=12)
        rt = p.add_run(desc)
        set_run(rt, size=12)

    h1("10.  Conclusion")
    body(
        "SentinelGPT demonstrates that integrating Agentic AI with a modern full-stack serverless "
        "architecture — React 19, FastAPI, SQLAlchemy, and Vercel — can fundamentally transform "
        "reactive, manual security operations into a proactive, fully automated cyber defense "
        "system. By replacing human-dependent log review with autonomous, continuously running "
        "software agents, the platform eliminates analyst alert fatigue, reduces incident response "
        "latency to near zero, and provides uninterrupted 24/7 network perimeter protection."
    )
    body(
        "The heuristic threat scoring engine, MITRE ATT&CK technique mapping, conversational AI "
        "triage assistant, and automated quarantine enforcement together establish a comprehensive "
        "and intelligent security operations platform that is scalable, transparent, and suitable "
        "for deployment in enterprise, financial, healthcare, and academic environments. The project "
        "sets a strong technical foundation for the next generation of AI-powered cybersecurity "
        "frameworks and demonstrates the practical value of agentic approaches in solving real-world "
        "security challenges at scale."
    )
    doc.add_page_break()

    # ═══════════════════════════════════════════════════════════════
    #  PAGE 9  —  REFERENCES + LIVE LINKS
    # ═══════════════════════════════════════════════════════════════
    h1("11.  References")
    refs = [
        "[1]  FastAPI Documentation — Available at: https://fastapi.tiangolo.com/",
        "[2]  React 19 Documentation — Available at: https://react.dev/",
        "[3]  MITRE ATT&CK Framework — Enterprise Tactics & Techniques. Available at: https://attack.mitre.org/",
        "[4]  SQLAlchemy ORM Documentation — Available at: https://docs.sqlalchemy.org/",
        "[5]  Vite Build Tool Documentation — Available at: https://vitejs.dev/",
        "[6]  Vercel Serverless Platform Documentation — Available at: https://vercel.com/docs",
        "[7]  Python 3.11 Documentation — Available at: https://docs.python.org/3/",
        "[8]  Uvicorn ASGI Server — Available at: https://www.uvicorn.org/",
        "[9]  MDN Web Docs: HTML5, CSS3, JavaScript — Available at: https://developer.mozilla.org/",
        "[10] Russell, S., & Norvig, P. (2021). Artificial Intelligence: A Modern Approach (4th ed.). Pearson.",
        "[11] Wooldridge, M. (2009). An Introduction to MultiAgent Systems (2nd ed.). John Wiley & Sons.",
        "[12] OWASP Top Ten Security Risks — Available at: https://owasp.org/Top10/",
    ]
    for ref in refs:
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after  = Pt(4)
        p.paragraph_format.line_spacing = Pt(18)
        r = p.add_run(ref)
        set_run(r, size=11)

    # Live Links section
    h1("12.  Live Deployment & GitHub Repository")
    body(
        "The SentinelGPT system is fully deployed and accessible online. Evaluators and reviewers "
        "may access the live working dashboard, interactive Swagger API documentation, and the complete "
        "project source code using the official links below:"
    )

    # Link table
    lt = doc.add_table(rows=3, cols=2)
    lt.style = "Table Grid"
    lt.alignment = WD_TABLE_ALIGNMENT.CENTER

    link_data = [
        ("Live SOC Dashboard", "https://sentinelgpt-ai.vercel.app"),
        ("Interactive Swagger API Docs", "https://sentinelgpt-ai.vercel.app/docs"),
        ("GitHub Source Repository", "https://github.com/Pravallika2025/sentigraud-ai-.git"),
    ]
    for i, (label, url) in enumerate(link_data):
        shade = (i % 2 == 0)
        cl, cr = lt.rows[i].cells
        alt_row(cl, shade); alt_row(cr, shade)
        pl = cl.paragraphs[0]
        pl.paragraph_format.space_before = Pt(4)
        pl.paragraph_format.space_after  = Pt(4)
        rl = pl.add_run(label); set_run(rl, bold=True, size=11)
        pu = cr.paragraphs[0]
        pu.paragraph_format.space_before = Pt(4)
        pu.paragraph_format.space_after  = Pt(4)
        ru = pu.add_run(url); set_run(ru, size=11)

    # Final note
    doc.add_paragraph()
    fn_p = doc.add_paragraph()
    fn_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    fn_p.paragraph_format.space_before = Pt(20)
    fn_r = fn_p.add_run(
        "─────────────────────────────────────────\n"
        "SentinelGPT  |  Pravallika Kalangi  |  24VV1F0044\n"
        "MCA 2nd Year  |  July 2026"
    )
    set_run(fn_r, italic=True, size=10)

    # ── SAVE ──────────────────────────────────────────────────────
    for path in save_paths:
        doc.save(path)
        print(f"Saved: {path}  ({round(os.path.getsize(path)/1024)} KB)")

# ─────────────────────────────────────────────────────────────────────────────
if __name__ == "__main__":
    build([OUT_FILE, DOCS_FILE])
    # Open the desktop copy immediately
    import subprocess
    subprocess.Popen(["start", "", OUT_FILE], shell=True)
    print("Opening document on screen...")
